import docx
import os
import re
from datetime import datetime
from .base_parser import BaseParser

class DocxParser(BaseParser):
    def parse(self):
        doc = docx.Document(self.filepath)
        filename = os.path.basename(self.filepath)
        
        # Extrair ID do nome do arquivo
        match = re.search(r'HistoriaUsuario_(\d+)', filename)
        hu_id = match.group(1) if match else "XXX"

        data = {
            "id": f"HU{hu_id}",
            "title": "Título não encontrado",
            "user_story": "",
            "revisions": [],
            "identification": {},
            "references": f"P:\\CGIA\\OSTENSIVO\\CGIT\\03_COIAI\\Lei do Bem\\GESTAO DE SISTEMAS\\ARTEFATOS\\{filename}",
            "generated_at": datetime.now().strftime("%d/%m/%Y %H:%M")
        }

        # 1. Revisões (Tabela 0)
        if len(doc.tables) > 0:
            table = doc.tables[0]
            for row in table.rows[1:]: # Pula header
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 4:
                    data["revisions"].append({
                        "version": cells[0],
                        "date": cells[1],
                        "description": cells[2].replace('\n', '<br>'),
                        "author": cells[3]
                    })

        # 2. Identificação (Tabela 1)
        if len(doc.tables) > 1:
            table = doc.tables[1]
            valid_keys = ["projeto", "requisitante", "gerente", "tema", "épico", "epico", "feature", "campo"]
            last_key = None
            
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if not cells: continue
                
                key_candidate = cells[0].lower()
                is_valid_key = any(vk in key_candidate for vk in valid_keys)
                
                key = None
                value = ""
                
                if len(cells) >= 2 and is_valid_key:
                    key = key_candidate
                    value = cells[1].replace('\n', '<br>')
                    last_key = key
                elif last_key: # Continuação
                    value = " ".join(cells).replace('\n', '<br>')
                    key = last_key
                    # Append logic simplificada para dicionário
                    current_val = data["identification"].get(self._normalize_key(key), "")
                    data["identification"][self._normalize_key(key)] = current_val + "<br>" + value if current_val else value
                    continue

                if key:
                    norm_key = self._normalize_key(key)
                    data["identification"][norm_key] = value

        # 3. Título (Tabela 2)
        if len(doc.tables) > 2:
            try:
                data["title"] = doc.tables[2].rows[0].cells[1].text.strip()
            except:
                pass

        # 4. User Story (Tabela 3)
        if len(doc.tables) > 3:
            table = doc.tables[3]
            capture = False
            for row in table.rows:
                text = row.cells[0].text.strip()
                if "Descrição" in text:
                    capture = True
                    continue
                if capture and text:
                    data["user_story"] = text
                    break
        
        return [data] # Retorna lista para compatibilidade com interface

    def _normalize_key(self, key):
        if "gerente" in key: return "gerente"
        if "projeto" in key: return "projeto"
        if "requisitante" in key: return "requisitante"
        if "tema" in key: return "tema"
        if "épico" in key or "epico" in key: return "epico"
        if "feature" in key: return "feature"
        return key
