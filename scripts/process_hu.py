import docx
import os
import glob
import re
import sys
import json
import subprocess
import time

# Configuration
PROJECT_NUMBER = 1 # 'rco-mcti/1' -> project number 1
OWNER = "rco-mcti"
REPO = "COIAI" # Assuming repo name from path, but better if implicit or env var. 
# For GH CLI, usually we run in the context of the repo.
GITHUB_PROJECT_ID = "rco-mcti/1"
ASSIGNEE = "rco-mcti"

class HuParser:
    def __init__(self, filepath):
        self.filepath = filepath
        self.filename = os.path.basename(filepath)
        self.doc = docx.Document(filepath)
        self.hu_id = self._extract_hu_id()
        self.data = {
            "revisions": [],
            "identification": {},
            "user_story": "",
            "title": "",
            "labels": []
        }

    def _extract_hu_id(self):
        match = re.search(r'HistoriaUsuario_(\d+)', self.filename)
        return match.group(1) if match else "XXX"

    def parse(self):
        self._parse_revisions()
        self._parse_identification()
        self._parse_title()
        self._parse_user_story()

    def _parse_revisions(self):
        if len(self.doc.tables) > 0:
            table = self.doc.tables[0]
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 4:
                    # Replace newlines in description with <br> to keep table format
                    description = cells[2].replace('\n', '<br>')
                    self.data["revisions"].append({
                        "version": cells[0],
                        "date": cells[1],
                        "description": description,
                        "author": cells[3]
                    })

    def _parse_identification(self):
        if len(self.doc.tables) > 1:
            table = self.doc.tables[1]
            last_key = None
            valid_keys = ["projeto", "requisitante", "gerente", "tema", "épico", "epico", "feature", "campo"]
            
            for row in table.rows:
                # Get text from cells, filter empty
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                
                if not cells:
                    continue
                
                key_candidate = cells[0].lower()
                is_valid_key = any(vk in key_candidate for vk in valid_keys)
                
                if len(cells) >= 2 and is_valid_key:
                    key = key_candidate
                    value = cells[1].replace('\n', '<br>')
                    last_key = key # Remember the key for continuation
                elif last_key:
                    # Treat as continuation
                    value = " ".join(cells).replace('\n', '<br>')
                    key = last_key
                    
                    # Append to existing value in data with <br>
                    if "gerente" in key: self.data["identification"]["gerente"] += "<br>" + value
                    elif "projeto" in key: self.data["identification"]["projeto"] += "<br>" + value
                    elif "requisitante" in key: self.data["identification"]["requisitante"] += "<br>" + value
                    elif "tema" in key: self.data["identification"]["tema"] += "<br>" + value
                    elif "épico" in key or "epico" in key: self.data["identification"]["epico"] += "<br>" + value
                    elif "feature" in key: self.data["identification"]["feature"] += "<br>" + value
                    continue # Skip the standard assignment below
                else:
                    continue

                if "gerente" in key: self.data["identification"]["gerente"] = value
                elif "projeto" in key: self.data["identification"]["projeto"] = value
                elif "requisitante" in key: self.data["identification"]["requisitante"] = value
                
                elif "tema" in key: 
                    self.data["identification"]["tema"] = value
                    self._extract_label(value)
                elif "épico" in key or "epico" in key: 
                    self.data["identification"]["epico"] = value
                    self._extract_label(value)
                elif "feature" in key: 
                    self.data["identification"]["feature"] = value
                    self._extract_label(value)

    def _extract_label(self, text):
        matches = re.findall(r'\b([A-Z]{1,2}\d{2,3})\b', text)
        for code in matches:
            if code not in self.data["labels"]:
                self.data["labels"].append(code)

    def _parse_title(self):
        if len(self.doc.tables) > 2:
            try:
                self.data["title"] = self.doc.tables[2].rows[0].cells[1].text.strip()
            except:
                self.data["title"] = "Título não encontrado"

    def _parse_user_story(self):
        if len(self.doc.tables) > 3:
            table = self.doc.tables[3]
            capture_next = False
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if cells and "Descrição" in cells[0]:
                    capture_next = True
                    continue
                if capture_next and cells:
                    self.data["user_story"] = cells[0]
                    break

    def generate_body(self):
        # Generates ONLY the body for gh issue create, NOT the frontmatter
        rev_rows = ""
        last_rev = self.data["revisions"][-1] if self.data["revisions"] else {"version": "", "date": "", "description": "", "author": ""}

        md = f"""
## 📝 Histórico de Revisões
| Campo         | Descrição |
| :------------ | :-------- |
| **Versão**    | {last_rev['version']} |
| **Data**      | {last_rev['date']} |
| **Descrição** | {last_rev['description']} |
| **Autor**     | {last_rev['author']} |

## 🆔 Identificação
| Campo                   | Descrição |
| :---------------------- | :-------- |
| **Projeto**             | {self.data['identification'].get('projeto', '')} |
| **Requisitante**        | {self.data['identification'].get('requisitante', '')} |
| **Gerente de Projetos** | {self.data['identification'].get('gerente', '')} |
| **Tema**                | {self.data['identification'].get('tema', '')} |
| **Épico**               | {self.data['identification'].get('epico', '')} |
| **Feature**             | {self.data['identification'].get('feature', '')} |


## 📖 Descrição (User Story)
{self.data['user_story']}

## 🔗 Referências
- P:\\CGIA\\OSTENSIVO\\CGIT\\03_COIAI\\Lei do Bem\\GESTAO DE SISTEMAS\\ARTEFATOS\\{self.filename}
"""
        return md

def run_command(command):
    try:
        result = subprocess.run(command, shell=True, check=True, capture_output=True, text=True)
        return result.stdout.strip()
    except subprocess.CalledProcessError as e:
        print(f"Error running command: {command}")
        print(e.stderr)
        return None

def issue_exists(title):
    # Check if issue exists by searching for the exact title
    # gh issue list --search "in:title [HU006]: ..." --json number --limit 1
    # We need to escape quote for shell
    safe_title = title.replace('"', '\\"')
    cmd = f'gh issue list --search "in:title \\"{safe_title}\\"" --json number --limit 1'
    output = run_command(cmd)
    if output:
        issues = json.loads(output)
        if issues:
            return issues[0]['number']
    return None

import argparse

def process_file(filepath, dry_run=False):
    try:
        print(f"--- Processando: {os.path.basename(filepath)} ---")
        hu_parser = HuParser(filepath)
        hu_parser.parse()
        
        title = f"[HU{hu_parser.hu_id}]: {hu_parser.data['title']}"
        
        body = hu_parser.generate_body()
        labels = ",".join(hu_parser.data['labels']) if hu_parser.data['labels'] else ""

        # Check existence only if not dry-run or if we want to simulate properly (requires GH CLI)
        existing_number = None
        if not dry_run:
            existing_number = issue_exists(title)
        
        if dry_run:
            print(f"🔍 [DRY-RUN] Título: {title}")
            print(f"🔍 [DRY-RUN] Labels: {labels}")
            print(f"🔍 [DRY-RUN] Assignee: {ASSIGNEE}")
            print(f"🔍 [DRY-RUN] Project: {GITHUB_PROJECT_ID}")
            # print(f"🔍 [DRY-RUN] Corpo da Issue:\n{body}") 
            # Commented out body print to reduce noise, unless debugging
            
            # Simulate update check if possible, or just print intent
            # For dry-run, we might not know if it exists unless we actually query (which is safe read-only)
            # Let's query even in dry-run to show what WOULD happen (Create vs Update)
            existing_number_dry = issue_exists(title)
            if existing_number_dry:
                    print(f"🔍 [DRY-RUN] Issue #{existing_number_dry} já existe. Seria ATUALIZADA.")
            else:
                    print(f"🔍 [DRY-RUN] Issue não existe. Seria CRIADA.")
            
            print("-" * 40)
            return

        if existing_number:
            print(f"🔄 Issue já existe: #{existing_number}. Atualizando...")
            update_cmd = [
                "gh", "issue", "edit", str(existing_number),
                "--title", title,
                "--body", body
            ]
            # if labels:
            #     update_cmd.extend(["--add-label", labels])
            
            try:
                subprocess.run(update_cmd, check=True, capture_output=True, text=True)
                print(f"✅ Issue #{existing_number} atualizada com sucesso.")
            except subprocess.CalledProcessError as e:
                print(f"❌ Erro ao atualizar issue #{existing_number}: {e}")
                if e.stderr: print(f"🔍 Detalhes: {e.stderr}")
            
            # Ensure it is in the project
            add_to_project(existing_number)
            return

        # Create Issue
        print(f"🚀 Criando issue: {title}")
        create_cmd = [
            "gh", "issue", "create",
            "--title", title,
            "--body", body,
            # "--assignee", ASSIGNEE, # Removed
            # "--project", GITHUB_PROJECT_ID # Removed to handle V2 separately
        ]
        
        # if labels:
        #     create_cmd.extend(["--label", labels])
            
        result = subprocess.run(create_cmd, check=True, capture_output=True, text=True)
        # Output of create is the URL, e.g. https://github.com/owner/repo/issues/123
        new_issue_url = result.stdout.strip()
        print(f"✅ Issue criada: {new_issue_url}")
        
        # Extract number from URL
        new_number = new_issue_url.split('/')[-1]
        add_to_project(new_number)
        
    except subprocess.CalledProcessError as e:
        print(f"❌ Erro ao processar {filepath}: {e}")
        if e.stderr:
            print(f"🔍 Detalhes do erro (stderr): {e.stderr}")
    except Exception as e:
        print(f"❌ Erro inesperado ao processar {filepath}: {e}")

def add_to_project(issue_number):
    print(f"🗂 Adicionando issue #{issue_number} ao projeto {PROJECT_NUMBER} do usuário {OWNER}...")
    try:
        # gh project item-add <project-number> --owner <owner> --url <issue-url>
        issue_url = f"https://github.com/{OWNER}/{REPO}/issues/{issue_number}"
        
        cmd = [
            "gh", "project", "item-add", str(PROJECT_NUMBER),
            "--owner", OWNER,
            "--url", issue_url
        ]
        
        subprocess.run(cmd, check=True, capture_output=True, text=True)
        print("✅ Adicionada ao projeto com sucesso.")
    except subprocess.CalledProcessError as e:
        print(f"⚠️ Aviso: Não foi possível adicionar ao projeto. Verifique permissões/ID. Erro: {e}")
        if e.stderr: print(f"🔍 Detalhes: {e.stderr}")

def main():
    parser = argparse.ArgumentParser(description="Process HU DOCX files and create GitHub Issues.")
    parser.add_argument("--dry-run", action="store_true", help="Run without creating issues on GitHub, printing output instead.")
    args = parser.parse_args()

    # Search recursively for .docx files in the current directory and subdirectories
    files = glob.glob("**/*.docx", recursive=True)
    
    # Filter out temporary/hidden files or system directories if needed
    files = [f for f in files if not any(part.startswith('.') for part in f.split(os.sep))]

    if not files:
        print("Nenhum arquivo .docx encontrado no repositório.")
        return

    print(f"Encontrados {len(files)} arquivos para processar.")

    for filepath in files:
        process_file(filepath, dry_run=args.dry_run)

if __name__ == "__main__":
    main()
